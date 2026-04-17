"""File utilities: extract raw text from PDF / DOCX resumes."""
import io
from typing import Union

import pdfplumber
import fitz  # PyMuPDF
from docx import Document

from app.utils.logger import logger


class UnsupportedFileError(Exception):
    """Raised when the uploaded file format is not supported."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber with PyMuPDF fallback."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        text = "\n".join(text_parts).strip()
        if text:
            return text
        logger.warning("pdfplumber returned empty text, falling back to PyMuPDF")
    except Exception as exc:
        logger.warning(f"pdfplumber failed: {exc}; falling back to PyMuPDF")

    # Fallback - PyMuPDF
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            text = "\n".join(page.get_text() for page in doc)
        return text.strip()
    except Exception as exc:
        logger.error(f"PyMuPDF also failed: {exc}")
        raise


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also include table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        logger.error(f"Failed to read DOCX: {exc}")
        raise


def extract_text(filename: str, file_bytes: Union[bytes, bytearray]) -> str:
    """Dispatch to the right parser based on file extension."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(bytes(file_bytes))
    if name.endswith(".docx"):
        return extract_text_from_docx(bytes(file_bytes))
    raise UnsupportedFileError(
        f"Unsupported file type for '{filename}'. Only PDF and DOCX are supported."
    )
